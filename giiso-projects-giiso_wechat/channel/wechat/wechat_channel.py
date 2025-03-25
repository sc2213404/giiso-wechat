import os
import re
import socket
import time
import xml.etree.ElementTree as ET
from queue import Empty
from threading import Thread

from select import select

from base.func_giiso import Giiso, create_user_folder
from channel.channel import Channel
from channel.common.singleton import singleton
from configuration import config
from logger import logger
from task import SalesMessageTask
from wcferry import Wcf, WxMsg

import random
import markdown
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

MAX_WECHAT_TEXT_LENGTH = 500  # 假设微信最大字数限制


__version__ = "39.2.4.0"
# max_messages=10
os.environ['ntwork_LOG'] = "ERROR"

def is_port_in_use(port):
    """检查端口是否被占用"""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        try:
            s.bind(("0.0.0.0", port))
            return False  # 端口未被占用
        except OSError:
            return True  # 端口被占用

def find_available_port(start_port):
    """找到一个可用的端口"""
    port = start_port
    while True:
        if not is_port_in_use(port) and not is_port_in_use(port + 1):
            return port
        port += 2  # 如果 port 或 port+1 被占用，跳到 port+2 重新开始检查
def markdown_to_docx(markdown_text, doc_path):
    """将 Markdown 文本转换为格式化的 Word 文档，支持1~6级标题并保留#符号"""
    doc = Document()
    
    lines = markdown_text.split("\n")
    for line in lines:
        if line.startswith("# "):  # 1级标题
            p = doc.add_paragraph(f"{line[2:]}")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
        elif line.startswith("## "):  # 2级标题
            p = doc.add_paragraph(f"{line[3:]}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        elif line.startswith("### "):  # 3级标题
            p = doc.add_paragraph(f"{line[4:]}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif line.startswith("#### "):  # 4级标题
            p = doc.add_paragraph(f"{line[5:]}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(11)
        elif line.startswith("##### "):  # 5级标题
            p = doc.add_paragraph(f"{line[6:]}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
        elif line.startswith("###### "):  # 6级标题
            p = doc.add_paragraph(f"{line[7:]}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
        elif line.strip() == "":  # 处理空行
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)  # 普通文本段落

    doc.save(doc_path)
@singleton
class WeChatChannel(Channel):
    NOT_SUPPORT_REPLYTYPE = []

    def __init__(self) -> None:
        super().__init__()
        self.wcf = None
        self.allContacts = None

    def startup(self):
        available_port = find_available_port(10086)
        logger.info(f"启动wcf, 端口：{available_port}")
        wcf = Wcf(port=available_port, debug=True)
        logger.info("微信已登录")
        self.wcf = wcf
        self.user_id = self.wcf.get_self_wxid()
        self.allContacts = self.getAllContacts()

        self.chat = Giiso(self.user_id)
        self.init_bot()

        # 机器人启动发送测试消息
        # robot.sendTextMsg("1", "filehelper")

        # 接收消息
        self.enableReceivingMsg()  # 加队列pi

        # 给机器人添加主动销售定时任务，每 5 分钟执行一次
        sales_message_task = SalesMessageTask(self.user_id)
        self.onEverySeconds(300, sales_message_task.run, self.send_text)

        logger.info(f"WeChatRobot【{__version__}】成功启动···")
        self.inited = True

        # 让机器人一直跑
        self.keepRunningAndBlockProcess()

    def get_user_info(self):
        if not self.wcf.is_login():
            logger.error("微信未登录，用户信息获取失败")
            return None
        return self.wcf.get_user_info()

    def cleanup(self) -> None:
        if not self.wcf:
            return
        self.wcf.cleanup()


    @staticmethod
    def value_check(args: dict) -> bool:
        if args:
            return all(value is not None for key, value in args.items() if key != 'proxy')
        return False

    def toAt(self, msg: WxMsg) -> bool:
        """处理被 @ 消息
        :param msg: 微信消息结构
        :return: 处理状态，`True` 成功，`False` 失败
        """
        return self.toChitchat(msg)

    def toChitchat(self, msg: WxMsg) -> bool:
        """闲聊，接入 ChatGPT
        """
        if not self.chat:  # 没接 ChatGPT，固定回复
            return False

        wxid = msg.sender if not msg.from_group() else msg.roomid
        wxname = self.wcf.getIdName(wxid)
        q = re.sub(r"@.*?[\u2005|\s]", "", msg.content).replace(" ", "")
        if q == "":
            q = "你在微信中被@了，请根据最近的消息进行回复"
        else:
            # 非空消息，如果当前智能体开启了文件回复，则优先匹配文件，匹配到文件则直接回复文件、不再进行知识库回复
            if self.fileReply == '1' and self.file_match(msg, q):
                return True

        rsp = self.chat.get_answer(self.send_image, q, wxid, wxname, msg.from_group())
        if not rsp:
            logger.error(f"无法从 ChatGPT 获得答案")
            return False

        if msg.from_group():
            self.sendTextMsg(rsp, msg.roomid, msg.sender)
        else:
            self.sendTextMsg(rsp, msg.sender)

        # self.wcf.send_rich_text(name='Giiso写作机器人',
        #                         account='gh_1743d764d201',
        #                         title='在奇妙大陆上经营你的动物商铺帝国——《萌不萌宠不宠》邀你打造商业传奇！',
        #                         digest='欢迎来到动物星球-《萌不萌宠不宠》，一个充满智慧动物与无限可能的奇幻世界！在这里，每个城市、每个国家都由不同的动物群体组成，它们有各自独特的文化和需求。',
        #                         url='https://mp.weixin.qq.com/s/TcCnInAQi1QEC5ebz22Wew',
        #                         thumburl='',
        #                         receiver=wxid)

        return True

    def file_match(self, msg: WxMsg, q: str) -> bool:
        wx_id = msg.sender if not msg.from_group() else msg.roomid
        wx_name = self.wcf.getIdName(wx_id)
        file_match_result = self.chat.get_file(q, wx_id, wx_name, msg.from_group())
        if not file_match_result:
            return False

        file_url = file_match_result['file_url']
        if not file_url:
            return False

        answer = file_match_result['answer']
        if answer:
            if msg.from_group():
                self.sendTextMsg(answer, msg.roomid, msg.sender)
            else:
                self.sendTextMsg(answer, msg.sender)

        self.wcf.send_image(path=file_url, receiver=msg.roomid if msg.from_group() else msg.sender)
        return True

    def processMsg(self, msg: WxMsg) -> None:
        """当接收到消息的时候，会调用本方法。如果不实现本方法，则打印原始消息。
        此处可进行自定义发送的内容,如通过 msg.content 关键字自动获取当前天气信息，并发送到对应的群组@发送者
        群号：msg.roomid  微信ID：msg.sender  消息内容：msg.content
        content = "xx天气信息为："
        receivers = msg.roomid
        self.sendTextMsg(content, receivers, msg.sender)
        """
        # logger.info(msg)
        if not WeChatChannel().inited:
            logger.error("WeChatChannel not inited")
            return
        # 群聊消息
        if msg.from_group():
            self.handle_group(msg)
        else:
            self.handle_single(msg)

    def handle_single(self, msg: WxMsg):
        return
        # 非群聊信息，按消息类型进行处理
        if msg.type == 37:  # 好友请求
            self.autoAcceptFriendRequest(msg)
            return

        # 系统信息
        if msg.type == 10000:
            self.sayHiToNewFriend(msg)
            return

        # 文本消息
        if msg.type == 0x01:
            # 让配置加载更灵活，自己可以更新配置。也可以利用定时任务更新。
            if msg.from_self():
                if msg.content == "^更新$":
                    config.reload()
                    logger.info("已更新")
            else:
                logger.info("触发闲聊")
                self.toChitchat(msg)  # 闲聊
            return

        # 图片消息
        if (msg.type == 0x03 or msg.type == 47) and (self.imageRec == '1'):
            temp_save_path = os.path.join(os.getcwd(), f'output/picture/{self.user_id}/{msg.sender}/')
            create_user_folder(temp_save_path)
            image_path = self.wcf.download_image(msg.id, msg.extra, dir=temp_save_path, timeout=30)
            rsp = self.chat.get_img_answer(image_path, msg.sender, self.wcf.getIdName(msg.sender))
            if rsp:
                self.sendTextMsg(rsp, msg.sender)
            return

        # 语音消息
        if msg.type == 34:
            temp_save_path = os.path.join(os.getcwd(), f'output/voice/{self.user_id}/{msg.sender}/')
            create_user_folder(temp_save_path)
            audio_path = self.wcf.get_audio_msg(msg.id, dir=temp_save_path, timeout=30)
            rsp = self.chat.get_voice_answer(audio_path, msg.sender, self.wcf.getIdName(msg.sender))
            if rsp:
                self.sendTextMsg(rsp, msg.sender)


    def handle_group(self, msg: WxMsg):
        if msg.roomid !="43741496000@chatroom":
            return
        else:
            # 非文本消息，忽略
            if msg.type != 0x01:
                return
            chat_turns_key = f"{msg.roomid}_{msg.sender}"
            self.message_count[chat_turns_key] = 0
            self.toAt(msg)
            return
            '''
            # 被@或者有关键词，都会触发闲聊，一旦触发本轮对话会回复10次
            chat_turns_key = f"{msg.roomid}_{msg.sender}"
            print(chat_turns_key)
            if msg.is_at(self.user_id):  # 被@
                self.message_count[chat_turns_key] = 0
                self.toAt(msg)
                return

            # 触发关键词
            if self.key_words and any(keyword in msg.content for keyword in self.key_words):
                self.message_count[chat_turns_key] = 0
                self.toAt(msg)
                return

            if chat_turns_key in self.message_count and self.message_count[chat_turns_key] < self.max_messages:
                self.message_count[chat_turns_key] += 1
                # 根据本次消息内容判断是否有对话意图
                q = re.sub(r"@.*?[\u2005|\s]", "", msg.content).replace(" ", "")
                rsp = self.chat.get_chat(query=q, reply_decisison=True, wxname=self.wcf.getIdName(msg.sender),
                                        is_group=False, receiver_wxid=msg.sender)
                if rsp.startswith('0'):
                    return
                else:
                    self.toAt(msg)
            '''

    def enableReceivingMsg(self) -> None:
        def innerProcessMsg(wcf: Wcf):
            while wcf.is_receiving_msg():
                try:
                    msg = wcf.get_msg()
                    logger.info(msg)
                    self.processMsg(msg)
                except Empty:
                    continue  # Empty message
                except Exception as e:
                    logger.error(f"Receiving message error: {e}")

        self.wcf.enable_receiving_msg()
        Thread(target=innerProcessMsg, name="GetMessage", args=(self.wcf,), daemon=True).start()

    def sendTextMsg(self, msg: str, receiver: str, at_list: str = "") -> None:
        """
        发送文本消息。
        - 若包含沟通性语句，则先单独发送沟通性部分。
        - 正文超长时，转换为 docx 发送。
        """
        communication_phrases = ["当然可以","可以","请注意", "需要确认", "请查看", "请尽快处理", "需要您的帮助"]
        
        # 解析沟通性语句和正文内容
        communication_text = "\n".join(
            [line for line in msg.split("\n") if any(phrase in line for phrase in communication_phrases)]
        )
        main_text = "\n".join(
            [line for line in msg.split("\n") if not any(phrase in line for phrase in communication_phrases)]
        )

        # 发送沟通性文本
        if communication_text:

            self.wcf.send_text(communication_text, receiver, at_list)

        # 处理正文内容
        if len(main_text) > MAX_WECHAT_TEXT_LENGTH:

            doc_path = f"{receiver}_message.docx"
            markdown_to_docx(main_text, doc_path)
            self.sendDocMsg(doc_path, receiver)
        elif main_text:

            self.wcf.send_text(main_text, receiver, at_list)
        # msg 中需要有 @ 名单中一样数量的 @
        ats = ""
        if at_list:
            if at_list == "notify@all":  # @所有人
                ats = " @所有人"
            else:
                wxids = at_list.split(",")
                for wxid in wxids:
                    # 根据 wxid 查找群昵称
                    ats += f" @{self.wcf.get_alias_in_chatroom(wxid, receiver)}"

        # 增加时间间隔，降低风险
        # if not is_send_now:
        #     sleep_time = len(msg) / random.randint(3, 6)
        #     time.sleep(sleep_time)

        # {msg}{ats} 表示要发送的消息内容后面紧跟@，例如 北京天气情况为：xxx @张三
        if ats == "":
            logger.info(f"To {receiver}: {msg}")
            self.wcf.send_text(f"{msg}", receiver, at_list)
        else:
            logger.info(f"To {receiver}: {ats}\r{msg}")
            self.wcf.send_text(f"{ats}\n\n{msg}", receiver, at_list)
    
    def sendDocMsg(self, msg: str, receiver: str) -> None:
        """
        当消息长度超出限制时，转换为 Markdown 格式并发送 DOCX 文档。
        :param msg: 要发送的消息文本
        :param receiver: 接收人 wxid 或者群 id
        """
        temp_dir = os.path.join(os.getcwd(), "temp")  # 确保临时目录路径正确
        os.makedirs(temp_dir, exist_ok=True)  # **确保目录存在**

        doc_path = os.path.join(temp_dir, "wechat_message.docx")
        # 转换 Markdown 为 docx
        # doc = Document()
        # html_text = markdown.markdown(msg)  # Markdown 转 HTML
        # doc.add_paragraph(html_text)  # 简单方式，或使用 python-docx 复杂处理
        # doc.save(doc_path)
        markdown_to_docx(msg, doc_path)

        # 发送文档
        success = self.send_image(doc_path, receiver)  # 这里用 send_image 发送 docx

        if success:
            logger.info(f"文档发送成功: {doc_path} -> {receiver}")
        else:
            logger.error(f"文档发送失败: {doc_path} -> {receiver}")

    def send_image(self, path: str, receiver: str) -> bool:
        status = self.wcf.send_image(path=path, receiver=receiver)
        return status == 0

    def send_text(self, text: str, receiver: str) -> bool:
        status = self.wcf.send_text(text, receiver)
        return status == 0
    

    def getAllContacts(self) -> dict:
        """
        获取联系人（包括好友、公众号、服务号、群成员……）
        格式: {"wxid": "NickName"}
        """
        contacts = self.wcf.query_sql("MicroMsg.db", "SELECT UserName, NickName FROM Contact;")
        return {contact["UserName"]: contact["NickName"] for contact in contacts}

    def autoAcceptFriendRequest(self, msg: WxMsg) -> None:
        try:
            xml = ET.fromstring(msg.content)
            v3 = xml.attrib["encryptusername"]
            v4 = xml.attrib["ticket"]
            scene = int(xml.attrib["scene"])
            self.wcf.accept_new_friend(v3, v4, scene)

        except Exception as e:
            logger.error(f"同意好友出错：{e}")

    def sayHiToNewFriend(self, msg: WxMsg) -> None:
        nickName = re.findall(r"你已添加了(.*)，现在可以开始聊天了。", msg.content)
        if nickName:
            # 添加了好友，更新好友列表
            self.allContacts[msg.sender] = nickName[0]
            self.sendTextMsg(f"Hi {nickName[0]}，我自动通过了你的好友请求。", msg.sender)




